
# 读取docx中的文本代码示例
import docx
import pandas as pd
import re

# 获取文档对象
file = docx.Document(r"C:\Users\S406-18\Desktop\突发事件的分类和分级.docx")
print("段落数:" + str(len(file.paragraphs)))

for para in file.paragraphs:
    print(para.text)
    # pattern0 = re.compile(r'）.*(信号)$')
    # tmp = pattern0.search(para.text)
    # list = []
    # if tmp:
    #     print(tmp.string)

# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)