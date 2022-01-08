import re
import docx
import pandas as pd
import csv


class doc2csc_extractor:
    def __init__(self):
        self.doc = docx.Document(r"C:\Users\S406-18\Desktop\气象灾害预警信号及防御指南.docx")
        self.doc0 = docx.Document(r"C:\Users\S406-18\Desktop\突发事件的分类和分级.docx")
        self.doc_paras = len(self.doc.paragraphs)
        self.disaster_warning_signals = []
        self.sub_disaster_warning_signal_standard = []
        self.sub_disaster_warning_signal_ekp = []

        self.disasters = []

    # 提取《气象灾害预警信号及防御指南》
    def extractor(self):
        # 灾害预警信号对应的行的位置
        signal_index = []
        for i in range(len(self.doc.paragraphs)):
            p = re.compile(r'(?<=、).*(信号)$')
            t = p.search(self.doc.paragraphs[i].text)
            if t:
                self.disaster_warning_signals.append({t.group(0).strip('\n'): []})
                signal_index.append({i: t.group(0).strip('\n')})

        sub_signal_index = []
        for i in range(len(signal_index)):
            if i == len(signal_index) - 1:
                next_i = len(self.doc.paragraphs)
            else:
                next_i = list(signal_index[i + 1])[0]
            j = list(signal_index[i])[0]
            while j < next_i:
                p = re.compile(r'(?<=）).*(信号)$')
                t = p.search(self.doc.paragraphs[j].text)
                if t:
                    # list(signal_index[i]) == 台风预警信号, t == 台风蓝色预警信号
                    list(self.disaster_warning_signals[i].values())[0].append(t.group(0).strip('\n'))
                    sub_signal_index.append({j: t.group(0).strip('\n')})
                j += 1

        for i in range(len(sub_signal_index)):
            if i == len(sub_signal_index) - 1:
                next_i = len(self.doc.paragraphs)
            else:
                next_i = list(sub_signal_index[i + 1])[0]
            j = list(sub_signal_index[i])[0]
            s = ''
            while j < next_i:
                p = re.compile(r'(标准)(:|：)(.*)')
                t = p.search(self.doc.paragraphs[j].text)
                if t:
                    s += t.group(3).strip('\n')
                    p0 = re.compile(r'防御')
                    k = j + 1
                    while not p0.search(self.doc.paragraphs[k].text):
                        s += self.doc.paragraphs[k].text.strip('\n')
                        k += 1
                j += 1
                if s != '':
                    self.sub_disaster_warning_signal_standard.append({list(sub_signal_index[i])[0]: s})
                    s = ''

        for i in range(len(sub_signal_index)):
            if i == len(sub_signal_index) - 1:
                next_i = len(self.doc.paragraphs)
            else:
                next_i = list(sub_signal_index[i + 1])[0]
            j = list(sub_signal_index[i])[0]
            s = ''
            while j < next_i:
                p = re.compile(r'防御')
                t = p.search(self.doc.paragraphs[j].text)
                if t:
                    p0 = re.compile(r'[1-9]\..*')
                    k = j + 1
                    while k != len(self.doc.paragraphs) and p0.search(self.doc.paragraphs[k].text):
                        s += self.doc.paragraphs[k].text + '\n'
                        k += 1
                        j = k - 1
                j += 1
                if s != '':
                    self.sub_disaster_warning_signal_ekp.append({list(sub_signal_index[i])[0]: s})
                    s = ''
        return signal_index, sub_signal_index

    # 提取《突发事件的分类和分级》
    def extractor0(self):
        disaster_index = []
        disaster_index_desc = []
        for i in range(len(self.doc0.paragraphs)):
            p = re.compile(r'(.    .)|(.  .  .)')
            t = p.search(self.doc0.paragraphs[i].text)
            if t:
                self.disasters.append({t.group(0).strip('\n').replace(' ', ''): []})
                disaster_index.append({i: t.group(0).strip('\n').replace(' ', '')})
                disaster_index_desc.append({i: self.doc0.paragraphs[i + 2].text + ''})

        index_sub_signal = []
        sub_signal_standard = []
        for i in range(len(disaster_index)):
            if i == len(disaster_index) - 1:
                next_i = len(self.doc0.paragraphs)
            else:
                next_i = list(disaster_index[i + 1])[0]
            j = list(disaster_index[i])[0]
            while j < next_i:
                p = re.compile(r'●　.*(信号)')
                t = p.search(self.doc0.paragraphs[j].text)
                if t:
                    index_sub_signal.append({j: t.group(0).strip('\n')[2:]})
                    sub_signal_standard.append({j: self.doc0.paragraphs[j + 1].text[3:]})
                j += 1

        disaster_ekp = []
        for i in range(len(disaster_index)):
            if i == len(disaster_index) - 1:
                next_i = len(self.doc0.paragraphs)
            else:
                next_i = list(disaster_index[i + 1])[0]
            j = list(disaster_index[i])[0]
            while j < next_i:
                p = re.compile(r'☆ .*')
                t = p.search(self.doc0.paragraphs[j].text)
                ekps = []
                while t:
                    ekps.append(t.group(0))
                    j += 1
                    k = j
                    t = p.search(self.doc0.paragraphs[k].text)
                if len(ekps) != 0:
                    disaster_ekp.append({list(disaster_index[i])[0]: ekps})
                ekps = []
                j += 1
        print('')

    def ext2csv(self):
        signal_index, sub_signal_index = self.extractor()
        rows = []
        for i in range(len(signal_index)):
            index = list(signal_index[i])[0]
            signal = list(signal_index[i].values())[0]
            if i == len(signal_index) - 1:
                next_index = self.doc_paras
            else:
                next_index = list(signal_index[i + 1])[0]
            row = []
            standard = self.sub_disaster_warning_signal_standard
            ekp = self.sub_disaster_warning_signal_ekp
            for j in range(len(sub_signal_index)):
                sub_index = list(sub_signal_index[j])[0]
                if (sub_index > index) and (sub_index < next_index):
                    row.append(signal)
                    row.append(index)
                    row.append(list(sub_signal_index[j].values())[0])
                    row.append(sub_index)
                    row.append(list(standard[j].values())[0])
                    row.append(list(ekp[j].values())[0])
                    rows.append(row)
                    row = []
        with open("test.csv", "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["signal", "signal_index", "sub_signal", "sub_index", "standard", "ekp"])
            writer.writerows(rows)
        # test = pd.DataFrame(rows, columns=["signal","signal_index","sub_signal", "standard", "ekp"])
        # test.to_csv('test.csv')
        print("done!")

    def read_csv(self):
        data = pd.read_csv('test.csv')
        print(data)


if __name__ == "__main__":
    extractor = doc2csc_extractor()
    extractor.ext2csv()
    extractor.read_csv()
    extractor.extractor0()
